# student_eval.py
# File comes in a pdf format. Will need to export a word file with 3 tables:
# - student
# - teacher_1
# - teacher_2
# - Columns: Reason, Score, Confidence, Alert
# will only work on mac

from docx import Document
# from docx.shared import Inches
from PyPDF2 import PdfFileReader
import sys
import os
import re
# f = open(os.path.expanduser("~/Desktop/somefile.txt"))

# texts = []


def extract_information(file):
    pdf_path = "~/Documents/{}.pdf".format(file)
    # print(("this is awesome file: %s"), pdf_path)
    data = []
    with open(os.path.expanduser(pdf_path), 'rb') as f:
        reader = PdfFileReader(f)
        NumPages = reader.getNumPages()

        all_text = ""

        for i in range(0, NumPages):
            page = reader.getPage(i)
            text = page.extractText()
            all_text += text

        # texts.extend([all_text])

        x = all_text.split("  ")
        length = len(x)
        for i in range(length):
            string = x[i]

            if 't score on' in string.lower():
                temp = re.findall('\d+', string)
                risk = has_risk(string)
                title = find_subtitle(string)
                # printer_helper(temp, risk, title)
                row = {'title': title, 'risk': risk,
                       't-score': temp[0], 'rank': temp[1]}
                data.append(row)

            elif 'composite scale t score' in string.lower():
                temp = re.findall('\d+', string)
                risk = has_risk(string)
                title = has_title(string)
                # printer_helper(temp, risk, title)
                row = {'title': title, 'risk': risk,
                       't-score': temp[0], 'rank': temp[-1]}
                data.append(row)

            else:
                pass
                # print("=========")
                # print(string)
                # print("=========")
    # print(all_text)
    f.close()
    return data


def has_risk(string):
    s = string.lower()
    risk = ""
    if 'at-risk classification' in s:
        risk = 'At-Risk'

    elif 'clinically' in s:
        risk = 'Clinically Significant'

    else:
        risk = 'None'

    return risk


def has_title(string):
    title = ""
    s = string.lower()
    if 'the externalizing problems composite' in s:
        title = "Externalizing Problems composite"

    elif 'the internalizing problems composite' in s:
        title = "Internalizing Problems composite"

    elif 'the school problems composite' in s:
        title = "Internalizing Problems composite"

    elif 'the behavioral symptoms index (bsi) composite' in s:
        title = "Behavioral Symptoms Index (BSI) composite"

    elif 'adaptive skills composite' in s:
        title = "Adaptive Skills composite"

    elif 'the internalizing problems composite' in s:
        title = "Internalizing Problems composite"

    return title


def find_subtitle(string):
    subtitle = ""
    split1 = string.split("is")
    s1 = split1[0]
    split2 = s1.split(" on ")
    subtitle += split2[-1]
    return subtitle


def printer_helper(temp, risk, title):
    print("\n")
    print("title: ", title)
    print("T score:", temp[0])
    print("percentile rank:", temp[-1])
    print("risk: ", risk)
    print("\n")


def create_table_doc(data, file):
    document = Document()

    document.add_heading('SCORES', 0)

    p = document.add_paragraph('Name: ')
    p.add_run('bold').bold = True

    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')

    data_size = len(data)

    table = document.add_table(rows=(1), cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Title'
    hdr_cells[1].text = 'T-Score'
    hdr_cells[2].text = 'Rank'
    hdr_cells[3].text = 'Risk'

    bold_rows = []

    for i in range(data_size):
        if 'composite' in data[i]['title'].lower():
            bold_rows.append(i+1)

        row_cells = table.add_row().cells
        row_cells[0].text = data[i]['title']
        row_cells[1].text = data[i]['t-score']
        row_cells[2].text = data[i]['rank']
        row_cells[3].text = data[i]['risk']

    print(bold_rows)

    # document.add_page_break()
    docx_path = "~/Documents/{}_results.docx".format(file)
    document.save(os.path.expanduser(docx_path))

# def make_row_bold(rows):
#     rows=len(rows)
#     for i in range(rows):
#         row = rows[i]
#         row.cells[0].paragraphs[0].add_run(reqdheaderList[i]).bold=True


#         def make_rows_bold(*rows):
#     for row in rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 for run in paragraph.runs:
#                     run.font.bold = True
if __name__ == '__main__':
    file = sys.argv[1]
    data = extract_information(file)
    create_table_doc(data, file)
    # print(data)
